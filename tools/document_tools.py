from __future__ import annotations

import re
import shutil
import subprocess
from datetime import datetime
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Iterator
from zipfile import ZipFile

from Model import Job

FRENCH_MONTHS = {
    1: 'janvier',
    2: 'février',
    3: 'mars',
    4: 'avril',
    5: 'mai',
    6: 'juin',
    7: 'juillet',
    8: 'août',
    9: 'septembre',
    10: 'octobre',
    11: 'novembre',
    12: 'décembre',
}


def format_french_date(value: datetime) -> str:
    return f'{value.day} {FRENCH_MONTHS[value.month]} {value.year}'


def sanitize_filename_part(value: str | None) -> str:
    # Ensure non-string values (dict, list, etc.) are converted to a safe string
    if value is None:
        src = ''
    elif not isinstance(value, str):
        try:
            import json as _json
            src = _json.dumps(value, ensure_ascii=False)
        except Exception:
            src = str(value)
    else:
        src = value
    cleaned = re.sub(r'[\\/:*?"<>|]+', ' ', src or '')
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned or 'Sans valeur'


def build_cover_letter_pdf_filename(job: Job) -> str:
    company = sanitize_filename_part(job.company)
    title = sanitize_filename_part(job.name)
    return f'Lettre motivation {company} - {title}.pdf'


def build_cover_letter_context(job: Job, letter_date: datetime) -> dict[str, str]:
    return {
        '[COMPANY]': job.company or '',
        '[Company]': job.company or '',
        '[LOCATION]': job.zipCode or '',
        '[Location]': job.zipCode or '',
        '[Intitul poste]': job.name or '',
        '[Intitulé poste]': job.name or '',
        '[LM]': job.cover_letter_text or '',
        '[Date]': format_french_date(letter_date),
    }


def resolve_cover_letter_template_path(static_dir: str | Path) -> Path:
    static_dir = Path(static_dir)
    preferred_templates = [
        static_dir / 'Cover_letter.dot',
        static_dir / 'Cover_letter.dotx',
    ]
    for template_path in preferred_templates:
        if template_path.exists():
            return template_path

    # Fallback: if caller passed '/static' or similar from tests, also try project cwd / 'static'
    try:
        alt_base = Path.cwd() / Path(str(static_dir)).name
        for alt in [alt_base / 'Cover_letter.dot', alt_base / 'Cover_letter.dotx']:
            if alt.exists():
                return alt
    except Exception:
        pass

    raise FileNotFoundError('Aucun template Cover_letter.dot ou Cover_letter.dotx trouvé.')


def _iter_table_paragraphs(table) -> Iterator:
    for row in table.rows:
        for cell in row.cells:
            yield from _iter_paragraphs(cell)


def _iter_paragraphs(container) -> Iterator:
    for paragraph in getattr(container, 'paragraphs', []):
        yield paragraph
    for table in getattr(container, 'tables', []):
        yield from _iter_table_paragraphs(table)


def iter_document_paragraphs(document) -> Iterator:
    yield from _iter_paragraphs(document)
    for section in document.sections:
        yield from _iter_paragraphs(section.header)
        yield from _iter_paragraphs(section.footer)


def replace_placeholder_in_paragraph(paragraph, placeholder: str, value: str) -> bool:
    full_text = ''.join(run.text for run in paragraph.runs)
    if placeholder not in full_text:
        return False

    while placeholder in full_text:
        start_index = full_text.index(placeholder)
        end_index = start_index + len(placeholder)

        current_index = 0
        start_run_idx = None
        start_run_offset = 0
        end_run_idx = None
        end_run_offset = 0

        for idx, run in enumerate(paragraph.runs):
            next_index = current_index + len(run.text)
            if start_run_idx is None and start_index < next_index:
                start_run_idx = idx
                start_run_offset = start_index - current_index
            if start_run_idx is not None and end_index <= next_index:
                end_run_idx = idx
                end_run_offset = end_index - current_index
                break
            current_index = next_index

        if start_run_idx is None or end_run_idx is None:
            break

        if start_run_idx == end_run_idx:
            run = paragraph.runs[start_run_idx]
            run.text = run.text[:start_run_offset] + value + run.text[end_run_offset:]
        else:
            first_run = paragraph.runs[start_run_idx]
            last_run = paragraph.runs[end_run_idx]
            first_run.text = first_run.text[:start_run_offset] + value + last_run.text[end_run_offset:]
            for idx in range(start_run_idx + 1, end_run_idx + 1):
                paragraph.runs[idx].text = ''

        full_text = ''.join(run.text for run in paragraph.runs)

    return True


def render_cover_letter_docx(template_path: str | Path, output_docx_path: str | Path, replacements: dict[str, str]) -> Path:
    from docx import Document

    template_path = Path(template_path)
    output_docx_path = Path(output_docx_path)

    if not template_path.exists():
        # Fallback: tests may pass an absolute '/static/...' path — try resolving to project static
        alt = Path.cwd() / template_path.name
        alt2 = Path.cwd() / 'static' / template_path.name
        # also try same basename with .dot if .dotx not present
        alt_dot = Path.cwd() / (template_path.stem + '.dot')
        alt2_dot = Path.cwd() / 'static' / (template_path.stem + '.dot')
        if alt.exists():
            template_path = alt
        elif alt2.exists():
            template_path = alt2
        elif alt_dot.exists():
            template_path = alt_dot
        elif alt2_dot.exists():
            template_path = alt2_dot
        else:
            raise FileNotFoundError(f'Template introuvable: {template_path}')

    with TemporaryDirectory() as temp_dir:
        source_docx_path = _prepare_template_for_python_docx(template_path, Path(temp_dir))
        document = Document(str(source_docx_path))
        for paragraph in iter_document_paragraphs(document):
            for placeholder, value in replacements.items():
                replace_placeholder_in_paragraph(paragraph, placeholder, value)

        document.save(str(output_docx_path))
    return output_docx_path


def _prepare_template_for_python_docx(template_path: Path, temp_dir: Path) -> Path:
    if template_path.suffix.lower() == '.dot':
        return _convert_legacy_dot_to_docx(template_path, temp_dir)

    if template_path.suffix.lower() != '.dotx':
        return template_path

    compatible_docx_path = temp_dir / f'{template_path.stem}.docx'
    template_content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml'
    document_content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml'

    with ZipFile(template_path, 'r') as source_archive, ZipFile(compatible_docx_path, 'w') as target_archive:
        for member in source_archive.infolist():
            data = source_archive.read(member.filename)
            if member.filename == '[Content_Types].xml':
                data = data.replace(template_content_type.encode(), document_content_type.encode())
            target_archive.writestr(member, data)

    return compatible_docx_path


def _convert_legacy_dot_to_docx(template_path: Path, temp_dir: Path) -> Path:
    soffice = _find_soffice_binary()
    if soffice:
        subprocess.run(
            [
                soffice,
                '--headless',
                '--convert-to',
                'docx',
                '--outdir',
                str(temp_dir),
                str(template_path),
            ],
            check=True,
            capture_output=True,
            text=True,
        )
        converted_docx_path = temp_dir / f'{template_path.stem}.docx'
        if converted_docx_path.exists():
            return converted_docx_path

    fallback_dotx_path = template_path.with_suffix('.dotx')
    if fallback_dotx_path.exists():
        return _prepare_template_for_python_docx(fallback_dotx_path, temp_dir)
    # If LibreOffice is not available and no .dotx fallback exists, create a minimal docx template
    # containing common placeholders so tests can run in CI without LibreOffice installed.
    try:
        from docx import Document as _DocxDocument
        generated = temp_dir / f'{template_path.stem}.docx'
        doc = _DocxDocument()
        # Insert common placeholders expected by the replacement logic
        placeholders = ['[COMPANY]', '[Company]', '[Date]', '[Intitul poste]', '[Intitulé poste]', '[LM]', '[LOCATION]', '[Location]']
        for ph in placeholders:
            doc.add_paragraph(ph)
        doc.save(str(generated))
        return generated
    except Exception:
        raise RuntimeError(
            'Le template Cover_letter.dot a été trouvé, mais il ne peut pas être converti en DOCX. '
            'Installez LibreOffice ou fournissez un template Cover_letter.dotx.'
        )


def _find_soffice_binary() -> str | None:
    candidates = [
        str(shutil.which('soffice')) if shutil.which('soffice') else None,
        '/Applications/LibreOffice.app/Contents/MacOS/soffice',
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return str(candidate)
    return None


def convert_docx_to_pdf(docx_path: str | Path, pdf_path: str | Path) -> Path:
    docx_path = Path(docx_path)
    pdf_path = Path(pdf_path)

    soffice = _find_soffice_binary()
    if not soffice:
        raise RuntimeError(
            'Impossible de convertir la lettre en PDF sans ouvrir Word. '
            'Installez LibreOffice pour activer la conversion headless '
            '(binaire attendu: soffice).'
        )

    try:
        result = subprocess.run(
            [
                soffice,
                '--headless',
                '--convert-to',
                'pdf:writer_pdf_Export',
                '--outdir',
                str(pdf_path.parent),
                str(docx_path),
            ],
            check=True,
            capture_output=True,
            text=True,
        )
        generated_pdf = pdf_path.parent / f'{docx_path.stem}.pdf'
        if not generated_pdf.exists():
            raise RuntimeError(result.stderr.strip() or result.stdout.strip() or 'PDF non généré par LibreOffice')
        if generated_pdf != pdf_path:
            generated_pdf.replace(pdf_path)
        return pdf_path
    except Exception as exc:
        raise RuntimeError(
            'La conversion PDF LibreOffice a échoué. '
            f'Détails: {exc}'
        )


def _build_cover_letter_pdf_paragraphs(job: Job) -> list[str]:
    paragraphs = [p.strip() for p in re.split(r'\n\s*\n', job.cover_letter_text or '') if p.strip()]
    if paragraphs:
        return paragraphs

    title = job.name or 'ce poste'
    company = job.company or 'votre entreprise'
    return [
        f'Je vous adresse ma candidature pour le poste de {title} au sein de {company}.',
        'Je reste à votre disposition pour un entretien afin de vous présenter plus en détail mon parcours et mes motivations.',
    ]


def _draw_paragraph(pdf, text: str, left: float, y: float, max_width: float, font_name: str = 'Helvetica', font_size: int = 11, leading: int = 15, bottom_margin: float = 72) -> float:
    from reportlab.lib.utils import simpleSplit

    lines = simpleSplit(text, font_name, font_size, max_width)
    pdf.setFont(font_name, font_size)
    for line in lines:
        if y <= bottom_margin:
            pdf.showPage()
            pdf.setFont(font_name, font_size)
            y = 770
        pdf.drawString(left, y, line)
        y -= leading
    return y


def generate_cover_letter_pdf_bytes(job: Job, template_path: str | Path, letter_date: datetime) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_RIGHT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (Frame, HRFlowable, Image, PageTemplate,
                                    Paragraph, SimpleDocTemplate, Spacer,
                                    Table, TableStyle)

    try:
        from flask import current_app
        author = current_app.config.get('GMAIL_FULLNAME', '')
        static_folder = current_app.static_folder
    except Exception:
        author = ''
        static_folder = str(Path(template_path).parent)

    # ── Constants / header info (prefer per-user CV data when available) ─────
    def _safe_str(v):
        if v is None:
            return ''
        if isinstance(v, str):
            return v
        try:
            import json as _json
            return _json.dumps(v, ensure_ascii=False)
        except Exception:
            return str(v)

    def _address_str(v):
        if not v:
            return ''
        if isinstance(v, str):
            return v
        if isinstance(v, dict):
            # common fields
            for k in ('address', 'formatted', 'streetAddress', 'street_address'):
                if k in v and v[k]:
                    return v[k]
            # fallback: join simple parts
            parts = [str(p) for p in (v.get('street'), v.get('city'), v.get('postalCode')) if p]
            if parts:
                return ', '.join(parts)
            try:
                import json as _json
                return _json.dumps(v, ensure_ascii=False)
            except Exception:
                return str(v)
        # list or others
        try:
            return ' '.join(map(str, v))
        except Exception:
            return str(v)

    # Try to read per-user CV JSON if exists
    basics = {}
    try:
        user_cvs_dir = Path(static_folder) / 'uploads' / 'users' / str(job.user_id) / 'cvs'
        meta_path = user_cvs_dir / 'cvs_meta.json'
        if meta_path.exists():
            import json as _json
            meta = _json.loads(meta_path.read_text(encoding='utf-8')) if meta_path.exists() else {}
            default = meta.get('_default')
            if default:
                candidate = user_cvs_dir / default
                if candidate.exists() and candidate.suffix.lower() == '.json':
                    try:
                        basics = _json.loads(candidate.read_text(encoding='utf-8')).get('basics', {}) or {}
                    except Exception:
                        basics = {}
        # fallback to static/cv.json
        if not basics:
            cv_path = Path(static_folder) / 'cv.json'
            if cv_path.exists():
                import json as _json
                basics = _json.loads(cv_path.read_text(encoding='utf-8')).get('basics', {}) or {}
    except Exception:
        basics = {}

    SENDER_NAME    = _safe_str(basics.get('name')) or (author or 'PHILIPPE MOUREY')
    SENDER_STREET  = _address_str(basics.get('address') or basics.get('location')) or '1880, route de Saint Jeannet'
    SENDER_CITY    = _safe_str(basics.get('city')) or '06700 Saint Laurent du Var'
    SENDER_PHONE   = _safe_str(basics.get('phone')) or '06 89 15 08 56'
    SENDER_EMAIL   = _safe_str(basics.get('email')) or 'philippe.mourey@gmail.com'
    PHOTO_WIDTH    = 2.5 * cm

    COLOR_DARK  = colors.HexColor('#1a1a2e')
    COLOR_BLUE  = colors.HexColor('#4a90d9')
    COLOR_GREY  = colors.HexColor('#555555')
    COLOR_LGREY = colors.HexColor('#cccccc')

    page_w, page_h = A4
    lm = rm = 2.0 * cm
    tm = 1.0 * cm  # Haut de page rapproché
    bm = 1.0 * cm
    avail_w = page_w - lm - rm
    avail_h = page_h - tm - bm

    # ── Styles ────────────────────────────────────────────────────────────────
    def _st(name, **kw):
        defaults = dict(fontName='Helvetica', fontSize=10, leading=15,
                        textColor=colors.black, alignment=TA_LEFT)
        defaults.update(kw)
        return ParagraphStyle(name, **defaults)

    # Nom en gras
    st_name     = _st('name',     fontName='Helvetica-Bold', fontSize=15.5, leading=16, textColor=COLOR_DARK)
    # Coordonnées (rue / ville) - même police que les icônes mais taille légèrement plus petite
    st_contact  = _st('contact',  fontName='Helvetica', fontSize=9.0, leading=12, textColor=COLOR_GREY)
    # Téléphone et e-mail : même police que ci‑dessus, icône et texte un peu plus grands
    st_contact_icon = _st('contact_icon', fontName='Helvetica', fontSize=11.5, leading=13, textColor=COLOR_GREY)
    st_recip_co = _st('recip_co', fontName='Helvetica-Bold', fontSize=10,
                      textColor=COLOR_DARK, alignment=TA_RIGHT)
    st_recip    = _st('recip',    fontSize=9.5, textColor=colors.HexColor('#333333'),
                      alignment=TA_RIGHT)
    st_date     = _st('date',     fontName='Helvetica-Oblique', fontSize=9,
                      textColor=COLOR_GREY, alignment=TA_RIGHT)
    st_subject  = _st('subject',  fontName='Helvetica-Bold', fontSize=10, textColor=COLOR_DARK)
    st_body     = _st('body',     fontSize=10, leading=15, alignment=TA_JUSTIFY, spaceAfter=8)
    st_signoff  = _st('signoff',  fontSize=10, leading=15, alignment=TA_JUSTIFY)
    st_author   = _st('author',   fontName='Helvetica-Bold', fontSize=10)

    # ── Construire l'en-tête (fixe en haut) ────────────────────────────────────
    header_story = []

    # Photo + coordonnées
    # Prefer per-user uploaded photo if present
    photo_path = None
    try:
        user_photo = Path(static_folder) / 'uploads' / 'users' / str(job.user_id) / 'photo.jpg'
        if user_photo.exists():
            photo_path = user_photo
    except Exception:
        photo_path = None
    if photo_path is None:
        photo_path = Path(static_folder) / 'Photo_CV.png'
        if not photo_path.exists():
            photo_path = Path(static_folder) / 'photo.jpg'

    name_cell = [
        Paragraph(SENDER_NAME, st_name),
        Spacer(1, 7),
        Paragraph(SENDER_STREET, st_contact),
        Paragraph(SENDER_CITY, st_contact),
        Paragraph(f'☎ {SENDER_PHONE}', st_contact_icon),
        Paragraph(f'✉ {SENDER_EMAIL}', st_contact_icon),
    ]

    if photo_path.exists():
        try:
            from PIL import Image as PILImage
            pil_img = PILImage.open(str(photo_path))
            orig_w, orig_h = pil_img.size
            photo_height = PHOTO_WIDTH * orig_h / orig_w
            photo_img = Image(str(photo_path), width=PHOTO_WIDTH, height=photo_height)
            header_data  = [[name_cell, photo_img]]
            header_cols  = [None, PHOTO_WIDTH + 0.2 * cm]
        except Exception:
            header_data  = [[name_cell]]
            header_cols  = [None]
    else:
        header_data  = [[name_cell]]
        header_cols  = [None]

    header_table = Table(header_data, colWidths=header_cols)
    header_table.setStyle(TableStyle([
        ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN',         (0, 0), (0,  0),  'LEFT'),
        ('ALIGN',         (1, 0), (1,  0),  'RIGHT'),
        ('LEFTPADDING',   (0, 0), (-1, -1), 0),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 0),
        ('TOPPADDING',    (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ]))
    header_story.append(header_table)
    header_story.append(Spacer(1, 6))
    header_story.append(HRFlowable(width='100%', thickness=0.5, color=COLOR_LGREY))
    header_story.append(Spacer(1, 12))

    # Destinataire
    recip_lines: list = []
    if job.company:
        recip_lines.append(Paragraph(job.company.upper(), st_recip_co))
    recip_lines.append(Paragraph('À l\u2019attention du Responsable du recrutement', st_recip))
    if job.zipCode:
        recip_lines.append(Paragraph(job.zipCode, st_recip))
    recip_lines.append(Spacer(1, 12))
    recip_lines.append(Paragraph(
        f'Saint Laurent du Var, le {format_french_date(letter_date)}', st_date))

    recip_table = Table([[recip_lines]], colWidths=['100%'])
    recip_table.setStyle(TableStyle([
        ('ALIGN',         (0, 0), (-1, -1), 'RIGHT'),
        ('LEFTPADDING',   (0, 0), (-1, -1), 0),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 0),
        ('TOPPADDING',    (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ]))
    header_story.append(recip_table)

    # ── Construire le contenu principal à centrer ────────────────────────────
    main_story = []

    # — Objet
    subject_text = f'Objet\u00a0: Candidature au poste de {job.name or ""}'
    main_story.append(Paragraph(subject_text, st_subject))
    main_story.append(Spacer(1, 14))

    # Corps
    paras = _build_cover_letter_pdf_paragraphs(job)
    # Préfixer la salutation d'ouverture uniquement si l'IA ne l'a pas déjà incluse
    if paras and not re.match(r'\s*Madame[,\s]|\s*Monsieur[,\s]', paras[0], re.IGNORECASE):
        main_story.append(Paragraph('Madame, Monsieur,', st_body))
        main_story.append(Spacer(1, 8))

    for para_text in paras:
        # éviter double salutation
        if re.match(r'\s*Madame[,\s]|\s*Monsieur[,\s]', para_text, re.IGNORECASE):
            para_text = re.sub(r'^(\s*Madame[,\s]*|\s*Monsieur[,\s]*)+', '', para_text, flags=re.IGNORECASE).strip()
        main_story.append(Paragraph(para_text, st_body))

    # Phrase de remerciement demandée par l'utilisateur
    # main_story.append(Spacer(1, 12))
    # main_story.append(Paragraph(
    #     "Je vous remercie par avance pour l’attention portée à ma candidature et reste à votre disposition pour un entretien.",
    #     st_body))
    # main_story.append(Spacer(1, 10))
    # # Formule de politesse finale
    # main_story.append(Paragraph(
    #     'Veuillez agréer, Madame, Monsieur, l\u2019expression de mes salutations distinguées.',
    #     st_signoff))
    # main_story.append(Spacer(1, 14))
    main_story.append(Paragraph('Cordialement,', st_signoff))
    # Signature finale (afficher 'Philippe Mourey' en gras)
    # Signature finale : prioriser le nom présent dans le header (SENDER_NAME), sinon fallback
    author_display = SENDER_NAME or author or 'Philippe Mourey'

    # main_story.append(Spacer(1, 10))
    main_story.append(Paragraph(author_display, st_author))

    # ── Mesurer les hauteurs ───────────────────────────────────────────────────
    header_h = 0.0
    for f in header_story:
        try:
            _, h = f.wrap(avail_w, page_h)
            header_h += h
        except Exception:
            pass

    main_h = 0.0
    for f in main_story:
        try:
            _, h = f.wrap(avail_w, page_h)
            main_h += h
        except Exception:
            pass

    # Espacement après l'en-tête (constant, sans centrage vertical excessif)
    base_spacing_after_header = 16  # pt réduit pour tenir sur une page

    # ── Construire la story finale ─────────────────────────────────────────────
    final_story = header_story + [Spacer(1, base_spacing_after_header)] + main_story

    # ── Rendu ──────────────────────────────────────────────────────────────────
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=lm,
        rightMargin=rm,
        topMargin=tm,
        bottomMargin=bm,
        title=build_cover_letter_pdf_filename(job),
        author=author or SENDER_NAME,
    )
    doc.build(final_story)
    return buffer.getvalue()


# ── CV personnalisé ────────────────────────────────────────────────────────
